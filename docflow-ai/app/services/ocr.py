"""
OCR Pipeline
============
Extracts raw text from:
  • Scanned PDFs  → pdf2image + pytesseract
  • Digital PDFs  → pdfminer (text layer first, fallback to OCR)
  • Images        → pytesseract
  • Word (.docx)  → python-docx
"""

import io
from pathlib import Path
from typing import Tuple

import pytesseract
from PIL import Image
from loguru import logger

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ─── Tesseract config ───────────────────────────────────────────────────────
# Page segmentation mode 6 = assume single uniform block of text
TESS_CONFIG = "--psm 6 --oem 3"


class OCRService:
    """Unified text extraction across file types."""

    SUPPORTED_TYPES = {
        ".pdf": "pdf",
        ".png": "image",
        ".jpg": "image",
        ".jpeg": "image",
        ".tiff": "image",
        ".tif": "image",
        ".bmp": "image",
        ".webp": "image",
        ".docx": "word",
        ".doc": "word",
    }

    def extract(self, file_path: Path) -> Tuple[str, str]:
        """
        Returns (raw_text, engine_used).
        engine_used is one of: 'tesseract', 'docx', 'pdf-text-layer'
        """
        suffix = file_path.suffix.lower()
        file_type = self.SUPPORTED_TYPES.get(suffix)

        if file_type is None:
            raise ValueError(f"Unsupported file type: {suffix}")

        if file_type == "pdf":
            return self._extract_pdf(file_path)
        elif file_type == "image":
            return self._extract_image(file_path)
        elif file_type == "word":
            return self._extract_docx(file_path)

        raise ValueError(f"Unknown type: {file_type}")

    # ── PDF ──────────────────────────────────────────────────────────────────

    def _extract_pdf(self, path: Path) -> Tuple[str, str]:
        """Try digital text layer first; fall back to OCR if needed."""
        text = self._try_pdf_text_layer(path)
        if text and len(text.strip()) > 100:
            logger.info(f"PDF text layer extracted: {len(text)} chars")
            return text, "pdf-text-layer"

        # Fallback: render pages → OCR
        logger.info(f"PDF text layer thin — falling back to Tesseract OCR")
        return self._pdf_ocr(path)

    def _try_pdf_text_layer(self, path: Path) -> str:
        """Extract embedded text using pdfminer (no OCR needed for digital PDFs)."""
        try:
            from pdfminer.high_level import extract_text
            return extract_text(str(path))
        except Exception as e:
            logger.warning(f"pdfminer failed: {e}")
            return ""

    def _pdf_ocr(self, path: Path) -> Tuple[str, str]:
        """Render PDF pages as images, run Tesseract on each."""
        if not PDF2IMAGE_AVAILABLE:
            raise RuntimeError("pdf2image not installed — cannot OCR scanned PDFs")

        pages = convert_from_path(str(path), dpi=300)
        texts = []
        for i, page_img in enumerate(pages):
            page_text = pytesseract.image_to_string(page_img, config=TESS_CONFIG)
            texts.append(page_text)
            logger.debug(f"OCR page {i+1}/{len(pages)}: {len(page_text)} chars")

        full_text = "\n\n".join(texts)
        logger.info(f"PDF OCR complete: {len(pages)} pages, {len(full_text)} chars")
        return full_text, "tesseract"

    # ── Image ─────────────────────────────────────────────────────────────────

    def _extract_image(self, path: Path) -> Tuple[str, str]:
        """Run Tesseract directly on image file."""
        img = Image.open(path)

        # Preprocess: convert to grayscale for better OCR accuracy
        img = img.convert("L")

        text = pytesseract.image_to_string(img, config=TESS_CONFIG)
        logger.info(f"Image OCR complete: {len(text)} chars")
        return text, "tesseract"

    # ── Word ──────────────────────────────────────────────────────────────────

    def _extract_docx(self, path: Path) -> Tuple[str, str]:
        """Extract text from .docx preserving paragraph structure."""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed")

        doc = Document(str(path))
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        text = "\n".join(parts)
        logger.info(f"DOCX extracted: {len(text)} chars, {len(doc.paragraphs)} paragraphs")
        return text, "docx"


ocr_service = OCRService()
